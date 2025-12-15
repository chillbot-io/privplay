"""Document Processor - Extract text from PDF, DOCX, clinical formats, and more.

Supports:
- Native PDF text extraction (pdfplumber)
- OCR for scanned PDFs (pytesseract + pdf2image) - optional
- DOCX extraction (python-docx)
- Excel (openpyxl, xlrd)
- RTF (striprtf)
- CSV/TSV
- Plain text files
- Images with OCR
- Clinical: HL7 v2, FHIR JSON, CCD/C-CDA XML

Usage:
    from privplay.documents import DocumentProcessor
    
    processor = DocumentProcessor()
    
    # Extract text from any supported format
    result = processor.extract("patient_records.pdf")
    result = processor.extract("lab_results.hl7")
    result = processor.extract("discharge_summary.xml")
    
    # Check capabilities
    print(processor.get_capabilities())
"""

import os
import csv
import json
import logging
from pathlib import Path
from typing import Tuple, Optional, Dict, Any, List, Union
from dataclasses import dataclass, field
from enum import Enum
from io import StringIO

logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    # Standard
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    RTF = "rtf"
    HTML = "html"
    CSV = "csv"
    TSV = "tsv"
    
    # Office
    XLSX = "xlsx"
    XLS = "xls"
    
    # Images
    PNG = "png"
    JPG = "jpg"
    JPEG = "jpeg"
    TIFF = "tiff"
    TIF = "tif"
    
    # Clinical
    HL7 = "hl7"
    FHIR_JSON = "fhir_json"
    CDA = "cda"
    CCD = "ccd"
    XML = "xml"  # Generic XML, may be CDA/CCD
    
    UNKNOWN = "unknown"


@dataclass
class ExtractionResult:
    """Result of document text extraction."""
    text: str
    doc_type: DocumentType
    page_count: int = 1
    metadata: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    ocr_used: bool = False
    
    # For structured data (HL7, FHIR, etc.)
    structured_data: Optional[Dict[str, Any]] = None
    
    @property
    def success(self) -> bool:
        return bool(self.text and not self.errors)


class DocumentProcessor:
    """Extract text from various document formats including clinical.
    
    Handles:
    - PDF (native text extraction, with optional OCR fallback)
    - DOCX/DOC (Microsoft Word)
    - XLSX/XLS (Excel)
    - RTF (Rich Text Format)
    - CSV/TSV (Comma/Tab Separated Values)
    - TXT (plain text)
    - Images (PNG, JPG, TIFF with OCR)
    - HL7 v2 (Lab results, ADT messages)
    - FHIR JSON (Modern healthcare format)
    - CCD/C-CDA/CDA XML (Clinical documents)
    """
    
    def __init__(
        self,
        enable_ocr: bool = True,
        ocr_language: str = "eng",
        min_text_length: int = 50,
    ):
        """Initialize document processor.
        
        Args:
            enable_ocr: Enable OCR for scanned PDFs and images
            ocr_language: Tesseract language code
            min_text_length: Minimum extracted text length before OCR fallback
        """
        self.enable_ocr = enable_ocr
        self.ocr_language = ocr_language
        self.min_text_length = min_text_length
        
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which extraction libraries are available."""
        self._deps = {}
        
        # PDF
        try:
            import pdfplumber
            self._deps['pdfplumber'] = True
        except ImportError:
            self._deps['pdfplumber'] = False
        
        try:
            from pypdf import PdfReader
            self._deps['pypdf'] = True
        except ImportError:
            self._deps['pypdf'] = False
        
        # Office
        try:
            import docx
            self._deps['docx'] = True
        except ImportError:
            self._deps['docx'] = False
        
        try:
            import openpyxl
            self._deps['openpyxl'] = True
        except ImportError:
            self._deps['openpyxl'] = False
        
        try:
            import xlrd
            self._deps['xlrd'] = True
        except ImportError:
            self._deps['xlrd'] = False
        
        try:
            from striprtf.striprtf import rtf_to_text
            self._deps['striprtf'] = True
        except ImportError:
            self._deps['striprtf'] = False
        
        # OCR
        if self.enable_ocr:
            try:
                import pytesseract
                from pdf2image import convert_from_path
                from PIL import Image
                self._deps['ocr'] = True
            except ImportError:
                self._deps['ocr'] = False
                logger.warning("OCR dependencies not available")
        else:
            self._deps['ocr'] = False
        
        # Clinical
        try:
            import hl7apy
            self._deps['hl7'] = True
        except ImportError:
            self._deps['hl7'] = False
        
        try:
            from fhir.resources import construct_fhir_element
            self._deps['fhir'] = True
        except ImportError:
            self._deps['fhir'] = False
        
        try:
            import lxml.etree
            self._deps['lxml'] = True
        except ImportError:
            self._deps['lxml'] = False
    
    def detect_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension and content."""
        path = Path(file_path)
        ext = path.suffix.lower().lstrip('.')
        
        type_map = {
            'pdf': DocumentType.PDF,
            'docx': DocumentType.DOCX,
            'doc': DocumentType.DOC,
            'txt': DocumentType.TXT,
            'text': DocumentType.TXT,
            'rtf': DocumentType.RTF,
            'html': DocumentType.HTML,
            'htm': DocumentType.HTML,
            'csv': DocumentType.CSV,
            'tsv': DocumentType.TSV,
            'xlsx': DocumentType.XLSX,
            'xls': DocumentType.XLS,
            'png': DocumentType.PNG,
            'jpg': DocumentType.JPG,
            'jpeg': DocumentType.JPEG,
            'tiff': DocumentType.TIFF,
            'tif': DocumentType.TIF,
            'hl7': DocumentType.HL7,
            'xml': DocumentType.XML,
            'json': DocumentType.FHIR_JSON,  # Assume FHIR, will validate
        }
        
        doc_type = type_map.get(ext, DocumentType.UNKNOWN)
        
        # For XML, try to detect if it's CDA/CCD
        if doc_type == DocumentType.XML:
            doc_type = self._detect_xml_type(path)
        
        return doc_type
    
    def _detect_xml_type(self, path: Path) -> DocumentType:
        """Detect specific XML clinical format."""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                header = f.read(2000)
            
            header_lower = header.lower()
            
            if 'clinicaldocument' in header_lower:
                if 'continuityofcaredocument' in header_lower:
                    return DocumentType.CCD
                return DocumentType.CDA
            
            return DocumentType.XML
        except:
            return DocumentType.XML
    
    def extract(self, file_path: str) -> ExtractionResult:
        """Extract text from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ExtractionResult with text and metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            return ExtractionResult(
                text="",
                doc_type=DocumentType.UNKNOWN,
                errors=[f"File not found: {file_path}"]
            )
        
        doc_type = self.detect_type(file_path)
        
        extractors = {
            DocumentType.PDF: self._extract_pdf,
            DocumentType.DOCX: self._extract_docx,
            DocumentType.DOC: self._extract_doc,
            DocumentType.TXT: self._extract_text,
            DocumentType.RTF: self._extract_rtf,
            DocumentType.HTML: self._extract_html,
            DocumentType.CSV: self._extract_csv,
            DocumentType.TSV: lambda p: self._extract_csv(p, delimiter='\t'),
            DocumentType.XLSX: self._extract_xlsx,
            DocumentType.XLS: self._extract_xls,
            DocumentType.PNG: self._extract_image,
            DocumentType.JPG: self._extract_image,
            DocumentType.JPEG: self._extract_image,
            DocumentType.TIFF: self._extract_image,
            DocumentType.TIF: self._extract_image,
            DocumentType.HL7: self._extract_hl7,
            DocumentType.FHIR_JSON: self._extract_fhir,
            DocumentType.CDA: self._extract_cda,
            DocumentType.CCD: self._extract_cda,  # Same parser
            DocumentType.XML: self._extract_xml,
        }
        
        extractor = extractors.get(doc_type)
        if extractor:
            return extractor(path)
        
        return ExtractionResult(
            text="",
            doc_type=doc_type,
            errors=[f"Unsupported document type: {doc_type.value}"]
        )
    
    # =========================================================================
    # PDF EXTRACTION
    # =========================================================================
    
    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """Extract text from PDF."""
        errors = []
        metadata = {}
        page_count = 0
        text = ""
        ocr_used = False
        
        if self._deps.get('pdfplumber'):
            try:
                text, page_count, metadata = self._extract_pdf_pdfplumber(path)
            except Exception as e:
                errors.append(f"pdfplumber failed: {e}")
        
        if not text and self._deps.get('pypdf'):
            try:
                text, page_count, metadata = self._extract_pdf_pypdf(path)
            except Exception as e:
                errors.append(f"pypdf failed: {e}")
        
        if len(text.strip()) < self.min_text_length and self._deps.get('ocr'):
            try:
                text, page_count = self._extract_pdf_ocr(path)
                ocr_used = True
                errors = []
            except Exception as e:
                errors.append(f"OCR failed: {e}")
        
        return ExtractionResult(
            text=text,
            doc_type=DocumentType.PDF,
            page_count=page_count,
            metadata=metadata,
            errors=errors,
            ocr_used=ocr_used,
        )
    
    def _extract_pdf_pdfplumber(self, path: Path) -> Tuple[str, int, Dict]:
        import pdfplumber
        
        text_parts = []
        metadata = {}
        
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            metadata = pdf.metadata or {}
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts), page_count, metadata
    
    def _extract_pdf_pypdf(self, path: Path) -> Tuple[str, int, Dict]:
        from pypdf import PdfReader
        
        reader = PdfReader(path)
        page_count = len(reader.pages)
        
        metadata = {}
        if reader.metadata:
            metadata = {
                'title': reader.metadata.title,
                'author': reader.metadata.author,
            }
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts), page_count, metadata
    
    def _extract_pdf_ocr(self, path: Path) -> Tuple[str, int]:
        import pytesseract
        from pdf2image import convert_from_path
        
        images = convert_from_path(path)
        page_count = len(images)
        
        text_parts = []
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image, lang=self.ocr_language)
            if page_text.strip():
                text_parts.append(f"--- Page {i+1} ---\n{page_text}")
        
        return "\n\n".join(text_parts), page_count
    
    # =========================================================================
    # OFFICE FORMATS
    # =========================================================================
    
    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extract text from DOCX."""
        if not self._deps.get('docx'):
            return ExtractionResult(
                text="",
                doc_type=DocumentType.DOCX,
                errors=["python-docx not installed"]
            )
        
        try:
            import docx
            
            doc = docx.Document(path)
            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                metadata = {
                    'title': props.title,
                    'author': props.author,
                }
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return ExtractionResult(
                text="\n\n".join(text_parts),
                doc_type=DocumentType.DOCX,
                metadata=metadata,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.DOCX,
                errors=[f"DOCX extraction failed: {e}"]
            )
    
    def _extract_doc(self, path: Path) -> ExtractionResult:
        """Extract text from legacy DOC format."""
        # Try antiword if available
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', str(path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return ExtractionResult(
                    text=result.stdout,
                    doc_type=DocumentType.DOC,
                )
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        return ExtractionResult(
            text="",
            doc_type=DocumentType.DOC,
            errors=["DOC format requires antiword: sudo apt install antiword"]
        )
    
    def _extract_xlsx(self, path: Path) -> ExtractionResult:
        """Extract text from XLSX."""
        if not self._deps.get('openpyxl'):
            return ExtractionResult(
                text="",
                doc_type=DocumentType.XLSX,
                errors=["openpyxl not installed"]
            )
        
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            text_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))
            
            wb.close()
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.XLSX,
                metadata={'sheet_count': len(wb.sheetnames)},
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.XLSX,
                errors=[f"XLSX extraction failed: {e}"]
            )
    
    def _extract_xls(self, path: Path) -> ExtractionResult:
        """Extract text from legacy XLS."""
        if not self._deps.get('xlrd'):
            return ExtractionResult(
                text="",
                doc_type=DocumentType.XLS,
                errors=["xlrd not installed"]
            )
        
        try:
            import xlrd
            
            wb = xlrd.open_workbook(path)
            text_parts = []
            
            for sheet in wb.sheets():
                text_parts.append(f"=== Sheet: {sheet.name} ===")
                
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    row_text = [str(cell) if cell else "" for cell in row]
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.XLS,
                metadata={'sheet_count': wb.nsheets},
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.XLS,
                errors=[f"XLS extraction failed: {e}"]
            )
    
    def _extract_rtf(self, path: Path) -> ExtractionResult:
        """Extract text from RTF."""
        if not self._deps.get('striprtf'):
            return ExtractionResult(
                text="",
                doc_type=DocumentType.RTF,
                errors=["striprtf not installed"]
            )
        
        try:
            from striprtf.striprtf import rtf_to_text
            
            rtf_content = path.read_text(encoding='utf-8', errors='ignore')
            text = rtf_to_text(rtf_content)
            
            return ExtractionResult(
                text=text,
                doc_type=DocumentType.RTF,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.RTF,
                errors=[f"RTF extraction failed: {e}"]
            )
    
    def _extract_csv(self, path: Path, delimiter: str = ',') -> ExtractionResult:
        """Extract text from CSV/TSV."""
        try:
            text_parts = []
            
            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f, delimiter=delimiter)
                for row in reader:
                    if any(row):
                        text_parts.append(" | ".join(row))
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.CSV if delimiter == ',' else DocumentType.TSV,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.CSV,
                errors=[f"CSV extraction failed: {e}"]
            )
    
    # =========================================================================
    # TEXT & HTML
    # =========================================================================
    
    def _extract_text(self, path: Path) -> ExtractionResult:
        """Extract text from plain text file."""
        try:
            try:
                text = path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                text = path.read_text(encoding='latin-1')
            
            return ExtractionResult(
                text=text,
                doc_type=DocumentType.TXT,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.TXT,
                errors=[f"Text extraction failed: {e}"]
            )
    
    def _extract_html(self, path: Path) -> ExtractionResult:
        """Extract text from HTML."""
        try:
            import re
            
            html = path.read_text(encoding='utf-8')
            
            html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.I)
            html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.I)
            text = re.sub(r'<[^>]+>', ' ', html)
            text = re.sub(r'\s+', ' ', text).strip()
            
            import html as html_module
            text = html_module.unescape(text)
            
            return ExtractionResult(
                text=text,
                doc_type=DocumentType.HTML,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.HTML,
                errors=[f"HTML extraction failed: {e}"]
            )
    
    # =========================================================================
    # IMAGES (OCR)
    # =========================================================================
    
    def _extract_image(self, path: Path) -> ExtractionResult:
        """Extract text from image using OCR."""
        if not self._deps.get('ocr'):
            return ExtractionResult(
                text="",
                doc_type=DocumentType.PNG,
                errors=["OCR not available (pytesseract, Pillow required)"]
            )
        
        try:
            import pytesseract
            from PIL import Image
            
            image = Image.open(path)
            text = pytesseract.image_to_string(image, lang=self.ocr_language)
            
            return ExtractionResult(
                text=text,
                doc_type=DocumentType(path.suffix.lower().lstrip('.')),
                ocr_used=True,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.PNG,
                errors=[f"Image OCR failed: {e}"]
            )
    
    # =========================================================================
    # CLINICAL FORMATS
    # =========================================================================
    
    def _extract_hl7(self, path: Path) -> ExtractionResult:
        """Extract text from HL7 v2 message."""
        if not self._deps.get('hl7'):
            # Fallback: basic parsing without hl7apy
            return self._extract_hl7_basic(path)
        
        try:
            from hl7apy.parser import parse_message
            from hl7apy.exceptions import HL7apyException
            
            content = path.read_text(encoding='utf-8', errors='ignore')
            # Normalize line endings
            content = content.replace('\r\n', '\r').replace('\n', '\r')
            
            try:
                msg = parse_message(content)
                text_parts = []
                structured = {}
                
                # Extract key segments
                if hasattr(msg, 'msh'):
                    msh = msg.msh
                    structured['message_type'] = str(msh.msh_9) if hasattr(msh, 'msh_9') else ''
                    structured['sending_app'] = str(msh.msh_3) if hasattr(msh, 'msh_3') else ''
                    text_parts.append(f"Message Type: {structured['message_type']}")
                
                # Patient info (PID segment)
                if hasattr(msg, 'pid'):
                    pid = msg.pid
                    patient_name = str(pid.pid_5) if hasattr(pid, 'pid_5') else ''
                    patient_id = str(pid.pid_3) if hasattr(pid, 'pid_3') else ''
                    dob = str(pid.pid_7) if hasattr(pid, 'pid_7') else ''
                    
                    text_parts.append(f"Patient: {patient_name}")
                    text_parts.append(f"Patient ID: {patient_id}")
                    text_parts.append(f"DOB: {dob}")
                    
                    structured['patient_name'] = patient_name
                    structured['patient_id'] = patient_id
                    structured['dob'] = dob
                
                # Observations (OBX segments)
                if hasattr(msg, 'obx'):
                    obx_list = msg.obx if isinstance(msg.obx, list) else [msg.obx]
                    text_parts.append("\nObservations:")
                    for obx in obx_list:
                        obs_id = str(obx.obx_3) if hasattr(obx, 'obx_3') else ''
                        obs_value = str(obx.obx_5) if hasattr(obx, 'obx_5') else ''
                        text_parts.append(f"  {obs_id}: {obs_value}")
                
                return ExtractionResult(
                    text="\n".join(text_parts),
                    doc_type=DocumentType.HL7,
                    structured_data=structured,
                    metadata={'parser': 'hl7apy'},
                )
                
            except HL7apyException:
                return self._extract_hl7_basic(path)
                
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.HL7,
                errors=[f"HL7 extraction failed: {e}"]
            )
    
    def _extract_hl7_basic(self, path: Path) -> ExtractionResult:
        """Basic HL7 parsing without hl7apy."""
        try:
            content = path.read_text(encoding='utf-8', errors='ignore')
            content = content.replace('\r\n', '\r').replace('\n', '\r')
            
            segments = content.split('\r')
            text_parts = []
            structured = {}
            
            for segment in segments:
                if not segment.strip():
                    continue
                
                fields = segment.split('|')
                seg_type = fields[0] if fields else ''
                
                if seg_type == 'MSH':
                    structured['message_type'] = fields[8] if len(fields) > 8 else ''
                    text_parts.append(f"Message Type: {structured['message_type']}")
                
                elif seg_type == 'PID':
                    patient_name = fields[5] if len(fields) > 5 else ''
                    patient_id = fields[3] if len(fields) > 3 else ''
                    dob = fields[7] if len(fields) > 7 else ''
                    
                    text_parts.append(f"Patient: {patient_name}")
                    text_parts.append(f"Patient ID: {patient_id}")
                    text_parts.append(f"DOB: {dob}")
                    
                    structured['patient_name'] = patient_name
                    structured['patient_id'] = patient_id
                    structured['dob'] = dob
                
                elif seg_type == 'OBX':
                    obs_id = fields[3] if len(fields) > 3 else ''
                    obs_value = fields[5] if len(fields) > 5 else ''
                    text_parts.append(f"Observation: {obs_id} = {obs_value}")
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.HL7,
                structured_data=structured,
                metadata={'parser': 'basic'},
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.HL7,
                errors=[f"HL7 basic extraction failed: {e}"]
            )
    
    def _extract_fhir(self, path: Path) -> ExtractionResult:
        """Extract text from FHIR JSON resource."""
        try:
            content = path.read_text(encoding='utf-8')
            data = json.loads(content)
            
            resource_type = data.get('resourceType', 'Unknown')
            text_parts = [f"FHIR Resource: {resource_type}"]
            
            # Extract based on resource type
            if resource_type == 'Patient':
                name = data.get('name', [{}])[0]
                given = ' '.join(name.get('given', []))
                family = name.get('family', '')
                text_parts.append(f"Patient: {given} {family}")
                
                if 'birthDate' in data:
                    text_parts.append(f"DOB: {data['birthDate']}")
                
                if 'identifier' in data:
                    for ident in data['identifier']:
                        text_parts.append(f"ID ({ident.get('system', 'unknown')}): {ident.get('value', '')}")
            
            elif resource_type == 'Observation':
                code = data.get('code', {}).get('text', data.get('code', {}).get('coding', [{}])[0].get('display', ''))
                text_parts.append(f"Observation: {code}")
                
                if 'valueQuantity' in data:
                    vq = data['valueQuantity']
                    text_parts.append(f"Value: {vq.get('value', '')} {vq.get('unit', '')}")
            
            elif resource_type == 'DiagnosticReport':
                code = data.get('code', {}).get('text', '')
                text_parts.append(f"Report: {code}")
                
                if 'conclusion' in data:
                    text_parts.append(f"Conclusion: {data['conclusion']}")
            
            elif resource_type == 'Bundle':
                entries = data.get('entry', [])
                text_parts.append(f"Bundle with {len(entries)} entries")
                for entry in entries[:10]:  # Limit for large bundles
                    resource = entry.get('resource', {})
                    text_parts.append(f"  - {resource.get('resourceType', 'Unknown')}")
            
            else:
                # Generic extraction
                text_parts.append(json.dumps(data, indent=2)[:2000])
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.FHIR_JSON,
                structured_data=data,
                metadata={'resource_type': resource_type},
            )
        except json.JSONDecodeError as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.FHIR_JSON,
                errors=[f"Invalid JSON: {e}"]
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.FHIR_JSON,
                errors=[f"FHIR extraction failed: {e}"]
            )
    
    def _extract_cda(self, path: Path) -> ExtractionResult:
        """Extract text from CDA/CCD/C-CDA XML."""
        if not self._deps.get('lxml'):
            return self._extract_xml(path)
        
        try:
            from lxml import etree
            
            tree = etree.parse(str(path))
            root = tree.getroot()
            
            # Handle namespaces
            ns = {'cda': 'urn:hl7-org:v3'}
            
            text_parts = []
            structured = {}
            
            # Document title
            title = root.find('.//cda:title', ns)
            if title is not None and title.text:
                text_parts.append(f"Document: {title.text}")
                structured['title'] = title.text
            
            # Patient name
            patient = root.find('.//cda:patient', ns)
            if patient is not None:
                name = patient.find('cda:name', ns)
                if name is not None:
                    given = name.find('cda:given', ns)
                    family = name.find('cda:family', ns)
                    patient_name = f"{given.text if given is not None else ''} {family.text if family is not None else ''}"
                    text_parts.append(f"Patient: {patient_name}")
                    structured['patient_name'] = patient_name
            
            # Extract all text content from sections
            for section in root.findall('.//cda:section', ns):
                section_title = section.find('cda:title', ns)
                if section_title is not None and section_title.text:
                    text_parts.append(f"\n=== {section_title.text} ===")
                
                # Get narrative text
                text_elem = section.find('cda:text', ns)
                if text_elem is not None:
                    section_text = etree.tostring(text_elem, method='text', encoding='unicode')
                    section_text = ' '.join(section_text.split())
                    if section_text:
                        text_parts.append(section_text)
            
            return ExtractionResult(
                text="\n".join(text_parts),
                doc_type=DocumentType.CDA,
                structured_data=structured,
                metadata={'format': 'CDA/C-CDA'},
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.CDA,
                errors=[f"CDA extraction failed: {e}"]
            )
    
    def _extract_xml(self, path: Path) -> ExtractionResult:
        """Generic XML text extraction."""
        try:
            if self._deps.get('lxml'):
                from lxml import etree
                tree = etree.parse(str(path))
                text = etree.tostring(tree, method='text', encoding='unicode')
            else:
                import xml.etree.ElementTree as ET
                tree = ET.parse(path)
                text = ET.tostring(tree.getroot(), method='text', encoding='unicode')
            
            text = ' '.join(text.split())
            
            return ExtractionResult(
                text=text,
                doc_type=DocumentType.XML,
            )
        except Exception as e:
            return ExtractionResult(
                text="",
                doc_type=DocumentType.XML,
                errors=[f"XML extraction failed: {e}"]
            )
    
    # =========================================================================
    # BATCH & UTILITIES
    # =========================================================================
    
    def extract_batch(self, file_paths: List[str]) -> List[ExtractionResult]:
        """Extract text from multiple documents."""
        return [self.extract(path) for path in file_paths]
    
    def extract_directory(
        self, 
        directory: str, 
        extensions: List[str] = None,
        recursive: bool = False,
    ) -> List[Tuple[str, ExtractionResult]]:
        """Extract text from all documents in a directory."""
        if extensions is None:
            extensions = ['.pdf', '.docx', '.txt', '.xlsx', '.hl7', '.xml', '.json']
        
        extensions = [ext.lower() if ext.startswith('.') else f'.{ext}' for ext in extensions]
        
        dir_path = Path(directory)
        results = []
        
        files = dir_path.rglob('*') if recursive else dir_path.glob('*')
        
        for file_path in files:
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                result = self.extract(str(file_path))
                results.append((str(file_path), result))
        
        return results
    
    def get_capabilities(self) -> Dict[str, bool]:
        """Get available extraction capabilities."""
        return {
            'pdf_native': self._deps.get('pdfplumber', False) or self._deps.get('pypdf', False),
            'pdf_ocr': self._deps.get('ocr', False),
            'docx': self._deps.get('docx', False),
            'xlsx': self._deps.get('openpyxl', False),
            'xls': self._deps.get('xlrd', False),
            'rtf': self._deps.get('striprtf', False),
            'csv': True,
            'txt': True,
            'html': True,
            'image_ocr': self._deps.get('ocr', False),
            'hl7': True,  # Basic parser always available
            'hl7_advanced': self._deps.get('hl7', False),
            'fhir': True,  # JSON parsing always available
            'cda_ccd': self._deps.get('lxml', False),
            'xml': True,
        }


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

_default_processor: Optional[DocumentProcessor] = None


def get_document_processor(**kwargs) -> DocumentProcessor:
    """Get or create default document processor."""
    global _default_processor
    if _default_processor is None:
        _default_processor = DocumentProcessor(**kwargs)
    return _default_processor


def extract_text(file_path: str) -> str:
    """Quick extraction - returns just the text or empty string."""
    processor = get_document_processor()
    result = processor.extract(file_path)
    return result.text if result.success else ""
